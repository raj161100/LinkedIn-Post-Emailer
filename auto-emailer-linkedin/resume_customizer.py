# resume_customizer.py
# Deterministic resume tailoring using python-docx (no external AI calls).
# Inserts a role-tailored summary & keyword highlight block at the top of a base resume.
# Optional: logs generated files in a tiny SQLite DB if db_path is provided.

import os
import re
import json
import hashlib
from datetime import datetime
from typing import Optional, List, Dict

from docx import Document

SAFE_MAX_LEN = 1200  # guardrail for extremely long job posts


def _extract_keywords(job_text: str) -> List[str]:
    """
    Very simple keyword extraction: scan for common tech terms & role hints.
    You can expand this list as needed.
    """
    canon = [
        # languages & runtimes
        "c#", ".net", "asp.net", "asp.net core", "java", "spring boot", "python", "node", "javascript", "typescript",
        # web & ui
        "react", "angular", "next.js", "blazor",
        # cloud
        "azure", "aws", "gcp", "cosmos", "dynamodb", "sql server", "postgres", "snowflake",
        # qa/sdet
        "selenium", "cypress", "playwright", "pytest", "junit",
        # data & infra
        "kafka", "spark", "airflow", "glue", "kinesis",
        # devops
        "docker", "kubernetes", "terraform", "github actions", "jenkins", "ci/cd",
        # patterns
        "microservices", "rest api", "event-driven", "serverless"
    ]
    text = job_text.lower()
    found = []
    for k in canon:
        if k in text:
            found.append(k)
    # de-dup preserving order
    seen = set()
    out = []
    for k in found:
        if k not in seen:
            out.append(k)
            seen.add(k)
    return out[:15]


def _digest(job_text: str, role_name: str) -> str:
    h = hashlib.sha1()
    h.update(role_name.encode("utf-8"))
    h.update(job_text[:SAFE_MAX_LEN].encode("utf-8"))
    return h.hexdigest()[:10]


def _insert_heading(paragraph, text: str):
    run = paragraph.add_run(text)
    run.bold = True


def generate_custom_resume(
    role_name: str,
    base_resume_path: str,
    job_text: str,
    output_dir: str = "assets/generated_resumes",
    company_hint: Optional[str] = None,
    job_link: Optional[str] = None,
    db_path: Optional[str] = None
) -> str:
    """
    Returns the path to a newly tailored .docx resume file.

    - role_name: e.g., ".NET Developer" or "SDET / QA Engineer"
    - base_resume_path: path to your base role resume template (.docx)
    - job_text: raw text from the LinkedIn post (used for tailoring)
    - company_hint: optional (derived from email domain or post)
    - job_link: optional (link to the job/post)
    - db_path: optional sqlite DB to record generated artifacts
    """
    if not os.path.exists(base_resume_path):
        raise FileNotFoundError(f"Base resume not found: {base_resume_path}")

    os.makedirs(output_dir, exist_ok=True)

    # Extract rough keywords and compute a unique filename suffix for this post
    kws = _extract_keywords(job_text or "")
    suffix = _digest(job_text or "", role_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"Resume_{role_name.replace(' ', '')}_{suffix}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)

    # Build a short tailored summary
    company_line = f" for {company_hint}" if company_hint else ""
    summary = (
        f"{role_name} focusing on scalable, production-grade systems{company_line}. "
        f"Skills mapped to this post: {', '.join(kws) if kws else 'core role competencies'}."
    )

    # Open base resume and insert a top section
    doc = Document(base_resume_path)

    # Insert our tailored block at the very top
    doc.paragraphs[0]._p.addprevious(doc.paragraphs[0]._p)  # ensure doc has a top
    p0 = doc.paragraphs[0]
    p0.clear()
    _insert_heading(p0, f"Target Role: {role_name}\n")

    p1 = doc.add_paragraph()
    p1.add_run(summary)

    if job_link:
        p2 = doc.add_paragraph()
        p2.add_run(f"Job Link: {job_link}")

    # Save the tailored resume
    doc.save(out_path)

    doc.core_properties.comments = (
    "Generated using RajAI Auto-Emailer — github.com/raj161100"
)


    # Optionally record in SQLite
    if db_path:
        try:
            import sqlite3
            os.makedirs(os.path.dirname(db_path), exist_ok=True)
            conn = sqlite3.connect(db_path)
            cur = conn.cursor()
            cur.execute("""
                CREATE TABLE IF NOT EXISTS generated_resumes(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    role TEXT,
                    path TEXT,
                    created_at TEXT,
                    keywords TEXT,
                    company TEXT,
                    job_link TEXT
                );
            """)
            cur.execute("""
                INSERT INTO generated_resumes(role, path, created_at, keywords, company, job_link)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                role_name, out_path, datetime.now().isoformat(timespec="seconds"),
                json.dumps(kws), company_hint, job_link
            ))
            conn.commit()
            conn.close()
        except Exception:
            # Non-fatal: if sqlite is missing/locked, just skip
            pass

    return out_path
